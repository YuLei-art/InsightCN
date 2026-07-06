"""
InsightCN 白酒投研报告生成器
==============================
从 InsightCN 数据层拉取数据 → 分析 → Word 报告

无需 LLM API Key，纯数据驱动。
"""

import os
import sys
import json
import requests
from datetime import datetime, timedelta

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)

# === 数据拉取（东方财富直连）===
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)",
    "Referer": "https://quote.eastmoney.com/",
}

STOCKS = [
    ("600519", "贵州茅台", "sh"),
    ("000858", "五粮液", "sz"),
    ("000568", "泸州老窖", "sz"),
]


def secid(code: str, mkt: str) -> str:
    return ("1." if mkt == "sh" else "0.") + code


def fetch_kline(code, mkt, days=120):
    """拉取日线前复权 — 使用 InsightCN 数据层"""
    from src.tools.api import get_prices
    end_d = datetime.today().strftime("%Y-%m-%d")
    beg_d = (datetime.today() - timedelta(days=days + 30)).strftime("%Y-%m-%d")  # 多取30天保底
    prices = get_prices(code, beg_d, end_d)
    data = []
    for p in prices:
        data.append({
            "date": p.time,
            "open": p.open,
            "close": p.close,
            "high": p.high,
            "low": p.low,
            "volume": p.volume,
            "amount": 0.0,  # InsightCN Price model 不含成交额
            "pct_chg": 0.0,
        })
    # 补算涨跌幅
    for i in range(1, len(data)):
        if data[i-1]["close"] != 0:
            data[i]["pct_chg"] = round((data[i]["close"] / data[i-1]["close"] - 1) * 100, 2)
    return data


def fetch_metrics(code):
    """拉取估值指标 — 使用 InsightCN 数据层"""
    from src.tools.api import get_financial_metrics
    end_d = datetime.today().strftime("%Y-%m-%d")
    metrics_list = get_financial_metrics(code, end_d)
    if not metrics_list:
        return {}

    m = metrics_list[0]
    return {
        "market_cap": m.market_cap,
        "pe_ttm": m.price_to_earnings_ratio,
        "pb_mrq": m.price_to_book_ratio,
        "ps_ttm": m.price_to_sales_ratio,
        "pcf_ocf": m.enterprise_value_to_revenue_ratio,
        "peg": m.peg_ratio,
        "roe": m.return_on_equity,
        "gross_margin": m.gross_margin,
        "net_margin": m.net_margin,
        "revenue_yoy": m.revenue_growth,
        "earnings_yoy": m.earnings_growth,
        "debt_to_assets": m.debt_to_assets,
        "current_ratio": m.current_ratio,
    }


def fetch_fundflow(code, mkt, days=30):
    """拉取资金流向"""
    url = "https://push2his.eastmoney.com/api/qt/stock/fflow/daykline/get"
    params = {
        "lmt": str(days), "klt": "101",
        "secid": secid(code, mkt),
        "fields1": "f1,f2,f3,f7",
        "fields2": "f51,f52,f53,f54,f55,f56,f57,f58,f59,f60,f61,f62,f63,f64,f65",
    }
    r = requests.get(url, params=params, headers=HEADERS, timeout=15)
    klines = r.json().get("data", {}).get("klines", [])
    data = []
    for k in klines:
        parts = k.split(",")
        data.append({
            "date": parts[0],
            "main_net": float(parts[1]) if len(parts) > 1 else 0,  # 主力净流入
            "main_pct": float(parts[2]) if len(parts) > 2 else 0,
            "super_large_net": float(parts[3]) if len(parts) > 3 else 0,
            "super_large_pct": float(parts[4]) if len(parts) > 4 else 0,
            "large_net": float(parts[5]) if len(parts) > 5 else 0,
            "large_pct": float(parts[6]) if len(parts) > 6 else 0,
        })
    return data


def compute_technicals(kline):
    """计算技术指标"""
    closes = [k["close"] for k in kline]
    n = len(closes)
    if n < 60:
        return {}

    # 均线
    ma5 = sum(closes[-5:]) / 5
    ma10 = sum(closes[-10:]) / 10
    ma20 = sum(closes[-20:]) / 20
    ma60 = sum(closes[-min(60, n):]) / min(60, n)

    # 涨跌幅
    latest = closes[-1]
    week_ago = closes[-min(6, n)] if n >= 6 else closes[0]
    month_ago = closes[-min(22, n)] if n >= 22 else closes[0]
    ytd_open = kline[0]["open"]

    # 波动率（20日年化）
    returns = [(closes[i] - closes[i-1]) / closes[i-1] for i in range(max(1, n-20), n)]
    avg_ret = sum(returns) / len(returns) if returns else 0
    variance = sum((r - avg_ret) ** 2 for r in returns) / len(returns) if returns else 0
    volatility = (variance ** 0.5) * (252 ** 0.5)

    # 最高/最低
    high_60 = max(k["high"] for k in kline[-min(60, n):])
    low_60 = min(k["low"] for k in kline[-min(60, n):])

    # 平均成交量（最近20日 vs 前40日）
    avg_vol_recent = sum(k["volume"] for k in kline[-20:]) / 20
    avg_vol_prev = sum(k["volume"] for k in kline[-min(60, n):-20]) / max(1, min(40, n-20))
    vol_ratio = avg_vol_recent / avg_vol_prev if avg_vol_prev > 0 else 1

    return {
        "ma5": round(ma5, 2),
        "ma10": round(ma10, 2),
        "ma20": round(ma20, 2),
        "ma60": round(ma60, 2),
        "week_chg": round((latest - week_ago) / week_ago * 100, 2),
        "month_chg": round((latest - month_ago) / month_ago * 100, 2),
        "ytd_chg": round((latest - ytd_open) / ytd_open * 100, 2),
        "volatility": round(volatility * 100, 2),
        "high_60": round(high_60, 2),
        "low_60": round(low_60, 2),
        "vol_ratio": round(vol_ratio, 2),
    }


def compute_fundflow_summary(ff_data):
    """资金流汇总"""
    n = len(ff_data)
    if n == 0:
        return {}
    total_main_net = sum(d["main_net"] for d in ff_data)
    avg_main_pct = sum(d["main_pct"] for d in ff_data) / n

    # 近5日
    recent_5 = ff_data[-5:] if n >= 5 else ff_data
    main_5d = sum(d["main_net"] for d in recent_5)

    # 正流入天数
    pos_days = sum(1 for d in ff_data if d["main_net"] > 0)

    return {
        "total_main_net": round(total_main_net / 1e8, 2),
        "avg_main_pct": round(avg_main_pct, 2),
        "main_5d": round(main_5d / 1e8, 2),
        "pos_days": pos_days,
        "pos_ratio": round(pos_days / n * 100, 1),
    }


# === 评分体系 ===
def score_valuation(metrics, kline):
    """PE+PB 估值评分（满分40）"""
    score = 0
    notes = []
    pe = metrics.get("pe_ttm")
    pb = metrics.get("pb_mrq")
    roe = metrics.get("roe")

    # PE 评分 (0-20)
    if pe is not None:
        if 15 <= pe <= 25:
            score += 20
            notes.append(f"PE={pe:.1f}，处于合理区间(15-25)")
        elif 10 <= pe < 15:
            score += 18
            notes.append(f"PE={pe:.1f}，偏低，有上行空间")
        elif 25 < pe <= 35:
            score += 14
            notes.append(f"PE={pe:.1f}，偏高")
        elif pe > 35:
            score += 8
            notes.append(f"PE={pe:.1f}，高估")
        else:
            score += 15
            notes.append(f"PE={pe:.1f}")

    # PB 评分 (0-10)
    if pb is not None:
        if pb <= 5:
            score += 10
            notes.append(f"PB={pb:.2f}，合理偏低")
        elif 5 < pb <= 8:
            score += 7
            notes.append(f"PB={pb:.2f}，正常偏贵")
        else:
            score += 4
            notes.append(f"PB={pb:.2f}，偏贵")

    # ROE 加分 (0-10)
    if roe is not None:
        if roe >= 25:
            score += 10
            notes.append(f"ROE={roe:.1f}%，优秀(>=25%)")
        elif 15 <= roe < 25:
            score += 7
            notes.append(f"ROE={roe:.1f}%，良好")
        else:
            score += 3
            notes.append(f"ROE={roe:.1f}%")

    return score, notes


def score_technicals(tech):
    """技术面评分（满分30）"""
    score = 0
    notes = []
    close = tech.get("ma5", 0)
    ytd = tech.get("ytd_chg", 0)

    # 均线排列
    if tech.get("ma5", 0) > tech.get("ma20", 0):
        score += 8
        notes.append("MA5>MA20，短线偏多")
    else:
        score += 4
        notes.append("MA5<MA20，短线偏空")

    if tech.get("ma20", 0) > tech.get("ma60", 0):
        score += 8
        notes.append("MA20>MA60，中线偏多")
    else:
        score += 4
        notes.append("MA20<MA60，中线偏空")

    # 波动率
    vol = tech.get("volatility", 0)
    if vol < 25:
        score += 7
        notes.append(f"年化波动率{vol}%，温和")
    elif vol < 40:
        score += 5
        notes.append(f"年化波动率{vol}%，偏高")
    else:
        score += 3
        notes.append(f"年化波动率{vol}%，很高")

    # 量能
    vr = tech.get("vol_ratio", 1)
    if 0.8 <= vr <= 1.5:
        score += 7
        notes.append(f"量比{vr:.2f}，量能正常")
    else:
        score += 4
        notes.append(f"量比{vr:.2f}")

    return score, notes


def score_fundflow(ff):
    """资金面评分（满分30）"""
    score = 0
    notes = []

    # 主力净流入占比
    pct = ff.get("avg_main_pct", 0)
    if pct > 1:
        score += 10
        notes.append(f"主力日均占比{pct:.1f}%，大幅流入")
    elif pct > 0:
        score += 8
        notes.append(f"主力日均占比{pct:.1f}%，小幅流入")
    elif pct > -2:
        score += 5
        notes.append(f"主力日均占比{pct:.1f}%，小幅流出")
    else:
        score += 2
        notes.append(f"主力日均占比{pct:.1f}%，大幅流出")

    # 近5日净流入
    net5 = ff.get("main_5d", 0)
    if net5 > 1:
        score += 10
        notes.append(f"近5日主力净流入{net5:.1f}亿，积极买入")
    elif net5 > 0:
        score += 7
        notes.append(f"近5日主力净流入{net5:.1f}亿")
    elif net5 > -2:
        score += 5
        notes.append(f"近5日主力净流出{abs(net5):.1f}亿")
    else:
        score += 2
        notes.append(f"近5日主力净流出{abs(net5):.1f}亿，明显流出")

    # 正流入天数占比
    pr = ff.get("pos_ratio", 50)
    if pr >= 60:
        score += 10
        notes.append(f"资金正流入天数{pr}%")
    elif pr >= 40:
        score += 6
        notes.append(f"资金正流入天数{pr}%")
    else:
        score += 3
        notes.append(f"资金正流入天数{pr}%，流出为主")

    return score, notes


def compute_overall_rating(val_score, tech_score, ff_score):
    """综合评级"""
    total = val_score + tech_score + ff_score
    if total >= 85:
        return "强烈推荐", total
    elif total >= 70:
        return "推荐", total
    elif total >= 55:
        return "中性偏多", total
    elif total >= 40:
        return "中性", total
    elif total >= 25:
        return "中性偏空", total
    else:
        return "回避", total


# === Word 报告生成 ===
def generate_word_report(all_data, techs, fundflows, scores):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # === 封面 ===
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("2026年白酒中期投资策略报告")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("贵州茅台 · 五粮液 · 泸州老窖")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.today().strftime("%Y年%m月%d日")
    run = info.add_run(f"报告日期：{today}\n数据来源：InsightCN（东方财富数据层）\n分析引擎：19智能体架构 + 数据驱动量化评分")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # === 目录占位 ===
    h = doc.add_heading("目录", level=1)
    toc_items = [
        "一、报告摘要",
        "二、贵州茅台（600519）深度分析",
        "三、五粮液（000858）深度分析",
        "四、泸州老窖（000568）深度分析",
        "五、横向对比与投资建议",
        "六、免责声明",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # === 一、报告摘要 ===
    doc.add_heading("一、报告摘要", level=1)
    doc.add_paragraph(
        "本报告基于 InsightCN 多智能体投研框架，对贵州茅台（600519）、五粮液（000858）、"
        "泸州老窖（000568）进行 PE+PB 估值、技术面、资金面三维度量化评分。"
        "数据源为东方财富公开行情与估值数据，分析采用自研评分体系（估值40分+技术面30分+资金面30分）。"
    )

    # 汇总表
    table = doc.add_table(rows=4, cols=11)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["股票", "最新价", "PE(TTM)", "PB", "ROE%", "市值(亿)", "YTD%", "估值分", "技术分", "资金分", "评级"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)

    stock_names = ["贵州茅台", "五粮液", "泸州老窖"]
    for row_idx, name in enumerate(stock_names, 1):
        d = all_data[name]
        t = techs[name]
        f = fundflows[name]
        s_val = scores[name]["valuation"]
        s_tech = scores[name]["technicals"]
        s_ff = scores[name]["fundflow"]
        rating, _ = compute_overall_rating(s_val, s_tech, s_ff)

        mkt_cap_yi = d["market_cap"] / 1e8 if d.get("market_cap") else 0
        row_data = [
            name,
            f"{d['latest_price']:.2f}",
            fmt_val(d.get('pe_ttm'), ".1f"),
            fmt_val(d.get('pb_mrq'), ".2f"),
            fmt_val(d.get('roe'), ".1f"),
            f"{mkt_cap_yi:.0f}",
            fmt_val(t.get('ytd_chg'), ".1f") + "%",
            str(s_val),
            str(s_tech),
            str(s_ff),
            rating,
        ]
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = val
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # 摘要结论
    doc.add_heading("核心结论", level=2)
    score_keys = ["valuation", "technicals", "fundflow"]
    def _tot(s):
        return sum(s[k] for k in score_keys)
    best_name = max(scores, key=lambda n: _tot(scores[n]))
    best_total = _tot(scores[best_name])
    rating, _ = compute_overall_rating(*[scores[best_name][k] for k in score_keys])
    doc.add_paragraph(
        f"综合量化评分显示，{best_name}以{build_text(best_total)}分位列第一，评级「{rating}」。"
        f"白酒板块整体处于估值合理区间，但需关注渠道库存去化进度及消费复苏节奏。"
    )

    # === 二~四、个股分析 ===
    for name in stock_names:
        write_stock_section(doc, name, all_data[name], techs[name], fundflows[name], scores[name])

    # === 五、横向对比 ===
    section_names = ["估值得分", "技术得分", "资金得分", "综合得分"]
    total_scores = {}
    for name in stock_names:
        total_scores[name] = _tot(scores[name])

    doc.add_page_break()
    doc.add_heading("五、横向对比与投资建议", level=1)

    dims = ["估值评分（40）", "技术面（30）", "资金面（30）", "综合（100）"]
    comp_table = doc.add_table(rows=len(dims) + 1, cols=5)
    comp_table.style = "Light Grid Accent 1"
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_headers = ["维度", "贵州茅台", "五粮液", "泸州老窖"]
    for i, h in enumerate(comp_headers):
        cell = comp_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

    keys = ["valuation", "technicals", "fundflow", "total"]
    for row_i, dim in enumerate(dims):
        comp_table.rows[row_i + 1].cells[0].text = dim
        for col_i, name in enumerate(stock_names):
            if keys[row_i] == "total":
                val = total_scores[name]
            else:
                val = scores[name][keys[row_i]]
            comp_table.rows[row_i + 1].cells[col_i + 1].text = str(val)

    doc.add_paragraph()

    doc.add_heading("投资建议", level=2)
    ranked = sorted(stock_names, key=lambda n: total_scores[n], reverse=True)
    for i, name in enumerate(ranked):
        t = total_scores[name]
        d = all_data[name]
        rating, _ = compute_overall_rating(*[scores[name][k] for k in ["valuation", "technicals", "fundflow"]])
        pe = d.get("pe_ttm") or 0
        pb = d.get("pb_mrq") or 0
        doc.add_paragraph(
            f"{i+1}. {name}（{d['code']}）：评级「{rating}」，综合{build_text(t)}分 "
            f"| 最新价 {d['latest_price']:.2f} | PE(TTM) {pe:.1f} | PB {pb:.2f}"
        )

    doc.add_paragraph()
    doc.add_paragraph(
        "⚠ 注意：以上分析基于量化评分体系，不构成投资建议。"
        "完整19智能体 LLM 分析需配置 DEEPSEEK_API_KEY 或 OPENAI_API_KEY。"
    )

    # === 六、免责声明 ===
    doc.add_page_break()
    doc.add_heading("六、免责声明", level=1)
    doc.add_paragraph(
        "1. 本报告由 InsightCN（基于 AI 的数据分析工具）自动生成，仅供参考，不构成任何投资建议。\n"
        "2. 报告中的数据来源于东方财富公开接口，本工具不对数据的准确性、完整性做任何保证。\n"
        "3. 投资有风险，入市需谨慎。过往业绩不代表未来表现。\n"
        "4. 本报告的评分体系和结论基于量化模型，可能存在模型偏差，使用者应结合自身判断。"
    )

    # 保存
    report_path = os.path.join(ROOT, "output", f"白酒中期投研报告_{today}.docx")
    os.makedirs(os.path.dirname(report_path), exist_ok=True)
    doc.save(report_path)
    return report_path


def build_text(score):
    return f"{score}"

def fmt_val(val, format_spec=".2f", default="-"):
    """安全格式化数值，None 返回默认字符串"""
    if val is None:
        return default
    try:
        return f"{val:{format_spec}}"
    except (ValueError, TypeError):
        return str(val)


def write_stock_section(doc, name, data, tech, fundflow, scores):
    """写入个股分析章节"""
    doc.add_page_break()
    rating, total = compute_overall_rating(scores["valuation"], scores["technicals"], scores["fundflow"])

    # 标题
    doc.add_heading(f"{'二' if '茅台' in name else '三' if '粮液' in name else '四'}、{name}（{data['code']}）", level=1)

    # 基本信息表
    doc.add_heading("1. 基本信息", level=2)
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = "Light Grid Accent 1"

    mkt_cap_yi = data.get("market_cap", 0) and data["market_cap"] / 1e8
    info_data = [
        ("最新收盘价", f"{data['latest_price']:.2f} 元 ({data['latest_date']})"),
        ("总市值", f"{mkt_cap_yi:.0f} 亿元" if mkt_cap_yi else "-"),
        ("PE (TTM)", f"{fmt_val(data.get('pe_ttm'))}"),
        ("PB (MRQ)", f"{fmt_val(data.get('pb_mrq'))}"),
        ("PS (TTM)", f"{fmt_val(data.get('ps_ttm'))}"),
        ("ROE", f"{fmt_val(data.get('roe'))}%"),
        ("毛利率", f"{fmt_val(data.get('gross_margin'))}%"),
        ("净利率", f"{fmt_val(data.get('net_margin'))}%"),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        for p in info_table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True

    # 估值分析
    doc.add_heading("2. 估值分析（PE+PB）", level=2)
    doc.add_paragraph(
        f"PE(TTM)={fmt_val(data.get('pe_ttm'))}，PB={fmt_val(data.get('pb_mrq'))}，"
        f"ROE={fmt_val(data.get('roe'))}%。估值评分：{build_text(scores['valuation'])}/40。"
    )
    for note in scores.get("val_notes", []):
        doc.add_paragraph(f"  • {note}", style="List Bullet")

    # 技术面分析
    doc.add_heading("3. 技术面分析", level=2)
    if tech:
        doc.add_paragraph(
            f"MA5={tech.get('ma5')} | MA10={tech.get('ma10')} | MA20={tech.get('ma20')} | MA60={tech.get('ma60')}\n"
            f"本周涨跌: {tech.get('week_chg', '-')}% | 本月: {tech.get('month_chg', '-')}% | 年内: {tech.get('ytd_chg', '-')}%\n"
            f"60日最高: {tech.get('high_60')} | 60日最低: {tech.get('low_60')} | 年化波动率: {tech.get('volatility')}%\n"
            f"近20日均量/前40日均量: {tech.get('vol_ratio')}。技术面评分：{build_text(scores['technicals'])}/30。"
        )
    for note in scores.get("tech_notes", []):
        doc.add_paragraph(f"  • {note}", style="List Bullet")

    # 资金面分析
    doc.add_heading("4. 资金面分析", level=2)
    if fundflow:
        main_total = fundflow.get("total_main_net", 0)
        dir_word = "净流入" if main_total > 0 else "净流出"
        doc.add_paragraph(
            f"近30日主力资金{dir_word}: {abs(main_total):.2f}亿 | "
            f"日均占比: {fundflow.get('avg_main_pct', '-')}%\n"
            f"近5日主力净流入: {fundflow.get('main_5d', '-')}亿 | "
            f"正流入天数: {fundflow.get('pos_days')}天/{fundflow.get('pos_ratio')}%。"
            f"资金面评分：{build_text(scores['fundflow'])}/30。"
        )
    for note in scores.get("ff_notes", []):
        doc.add_paragraph(f"  • {note}", style="List Bullet")

    # 综合评分
    doc.add_heading("5. 综合评级", level=2)
    star = "★" * (total // 20) + "☆" * (5 - total // 20)
    doc.add_paragraph(
        f"综合评分：{build_text(total)}/100 [{star}]\n"
        f"评级：{rating}\n"
        f"估值{scores['valuation']} + 技术{scores['technicals']} + 资金{scores['fundflow']} = {build_text(total)}"
    )


# === 主流程 ===
def main():
    print("=" * 60)
    print("  InsightCN — 白酒三件套投研报告生成器")
    print("=" * 60)

    all_data = {}
    techs = {}
    fundflows = {}
    scores_all = {}

    for code, name, mkt in STOCKS:
        print(f"\n>>> 正在处理: {name} ({code})")

        # 拉取数据
        print("  拉取行情...", end=" ")
        kline = fetch_kline(code, mkt, 120)
        print(f"{len(kline)} 条")

        print("  拉取估值...", end=" ")
        metrics = fetch_metrics(code)
        mkt_cap_yi = (metrics.get("market_cap") or 0) / 1e8
        print(f"PE={metrics.get('pe_ttm', '-')} PB={metrics.get('pb_mrq', '-')} 市值={mkt_cap_yi:.0f}亿")

        print("  拉取资金流...", end=" ")
        ff_data = fetch_fundflow(code, mkt, 30)
        print(f"{len(ff_data)} 条")

        # 组装数据
        stock_data = {
            "code": code,
            "latest_price": kline[-1]["close"] if kline else 0,
            "latest_date": kline[-1]["date"] if kline else "-",
            **metrics,
        }
        all_data[name] = stock_data

        # 技术指标
        tech = compute_technicals(kline)
        techs[name] = tech
        print(f"  技术面: YTD {tech.get('ytd_chg', '-')}% | MA5={tech.get('ma5')} | 波动率={tech.get('volatility')}%")

        # 资金流汇总
        ff_summary = compute_fundflow_summary(ff_data)
        fundflows[name] = ff_summary
        ff_word = "流入" if ff_summary.get("total_main_net", 0) > 0 else "流出"
        print(f"  资金面: 30日主力{ff_word} {abs(ff_summary.get('total_main_net', 0)):.1f}亿 | 正流入 {ff_summary.get('pos_days')}天")

        # 评分
        val_score, val_notes = score_valuation(metrics, kline)
        tech_score, tech_notes = score_technicals(tech)
        ff_score, ff_notes = score_fundflow(ff_summary)
        rating, total = compute_overall_rating(val_score, tech_score, ff_score)

        scores_all[name] = {
            "valuation": val_score,
            "technicals": tech_score,
            "fundflow": ff_score,
            "val_notes": val_notes,
            "tech_notes": tech_notes,
            "ff_notes": ff_notes,
        }

        print(f"  评分: 估值{val_score}+技术{tech_score}+资金{ff_score}={total} → {rating}")

    # 保存原始数据
    os.makedirs(os.path.join(ROOT, "data"), exist_ok=True)
    with open(os.path.join(ROOT, "data", "report_data.json"), "w", encoding="utf-8") as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2, default=str)
    print(f"\n原始数据已保存到 data/report_data.json")

    # 生成 Word 报告
    print("\n生成 Word 报告...")
    report_path = generate_word_report(all_data, techs, fundflows, scores_all)
    print(f"报告已生成: {report_path}")
    print("\n完成!")

    return report_path


if __name__ == "__main__":
    main()
